import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import tempfile
import zipfile
import io
from datetime import datetime

# Configurações do rodapé
RODAPE_CONFIG = {
    'endereco': 'Avenida Cristovão Colombo, nº 485, 4º andar, Savassi, Belo Horizonte/MG',
    'telefone': '(31) 9 9703-9242',
    'email': 'contato@icaadvocacia.com.br',
    'cor_fundo': (0, 0, 102),  # RGB para azul escuro (navy)
    'cor_texto': (255, 255, 255),  # RGB para branco
    'largura': '150%',  # Define a largura do rodapé como 100% da página
    'altura': '150%'
}

# Configurações de formatação
FORMATO_CONFIG = {
    'fonte_padrao': 'Arial',
    'tamanho_fonte_normal': 12,
    'tamanho_fonte_titulo': 12,
    'cor_titulo': (59, 75, 160),  # RGB para azul ICA (#3B4BA0)
    'cor_secao': (59, 75, 160),  # RGB para azul ICA
    'cor_linha': (192, 192, 192),  # RGB para cinza claro
    'espacamento_antes': Pt(6),
    'espacamento_depois': Pt(6),
    'espacamento_linha': 1.5
}

def criar_cabecalho(doc, logo_path=None):
    """
    Cria cabeçalho com logo ICA centralizado.
    """
    section = doc.sections[0]
    header = section.header

    # Limpar cabeçalho existente
    for para in header.paragraphs:
        para.clear()

    # Criar parágrafo para o logo
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adicionar logo se fornecido
    if logo_path and os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(2.5))  # Largura de 2.5 polegadas

    # Espaçamento após o logo
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.space_before = Pt(12)


def criar_rodape(doc, config):
    """
    Cria rodapé personalizado com fundo colorido e informações do escritório.
    """
    section = doc.sections[0]
    footer = section.footer

    # Salvar as margens originais do documento
    original_left_margin = section.left_margin
    original_right_margin = section.right_margin

    # Limpar rodapé existente
    for para in footer.paragraphs:
        para.clear()

    # Criar parágrafo do rodapé
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adicionar fundo colorido
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '{:02x}{:02x}{:02x}'.format(*config['cor_fundo']))
    p._element.get_or_add_pPr().append(shading_elm)

    # Configurar altura do parágrafo para aproximadamente 2.4 cm
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(50)

    # Adicionar espaço antes do texto para centralizá-lo verticalmente
    run_space_before = p.add_run("\n\n")  # Adiciona espaço no início
    run_space_before.font.size = Pt(2)  # Tamanho menor para controle fino

    # Linha 1: Endereço
    run1 = p.add_run(config['endereco'])
    run1.font.color.rgb = RGBColor(*config['cor_texto'])
    run1.font.size = Pt(10)
    run1.font.name = 'Arial'

    # Linha 2: Telefone e email
    p.add_run('\n')
    run2 = p.add_run(f"{config['telefone']} | {config['email']}")
    run2.font.color.rgb = RGBColor(*config['cor_texto'])
    run2.font.size = Pt(10)
    run2.font.name = 'Arial'
    p.add_run("\n\n\n")  # Add extra lines at the bottom


def adicionar_linha_horizontal(paragrafo, cor_rgb=(192, 192, 192)):
    p = paragrafo._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # Tamanho da linha
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '{:02x}{:02x}{:02x}'.format(*cor_rgb))
    pBdr.append(bottom)
    pPr.append(pBdr)


def aplicar_formatacao_paragrafo(paragrafo, alinhamento='justify', negrito=False,
                                 italico=False, tamanho_fonte=12, espacamento_antes=6,
                                 espacamento_depois=6, espacamento_linha=1.5, 
                                 cor_texto=None, recuo_lista=False, recuo_primeira_linha=True):
    # Alinhamento
    if alinhamento == 'center':
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alinhamento == 'justify':
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alinhamento == 'left':
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Espaçamento
    paragrafo.paragraph_format.space_before = Pt(espacamento_antes)
    paragrafo.paragraph_format.space_after = Pt(espacamento_depois)
    paragrafo.paragraph_format.line_spacing = espacamento_linha

    # Recuo para itens de lista ou primeira linha
    if recuo_lista:
        paragrafo.paragraph_format.left_indent = Inches(0.25)
        paragrafo.paragraph_format.first_line_indent = Inches(-0.25)
    elif recuo_primeira_linha:
        paragrafo.paragraph_format.first_line_indent = Cm(1.27)  # 0.5 polegadas

    # Formatação de fonte
    for run in paragrafo.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(tamanho_fonte)
        run.bold = negrito
        run.italic = italico
        if cor_texto:
            run.font.color.rgb = RGBColor(*cor_texto)


def detectar_tipo_paragrafo(texto, em_pedidos=False):
    """
    Detecta o tipo de parágrafo com base em características específicas.
    Implementa uma abordagem escalável com verificações em ordem de prioridade.
    """
    texto_limpo = texto.strip()

    # ETAPA 1: Verificações de estrutura específica (maior prioridade)
    
    # Detecção da seção de Pedidos
     if re.match(r'^(PEDIDOS|POR TUDO ISSO|DOS PEDIDOS|DO PEDIDO|IV[\s]*[.\-–—]+[\s]*DOS PEDIDOS)', texto_limpo, re.IGNORECASE):
        return 'secao_pedidos', True, 'center'  # Título centralizado

    # Se já estamos na seção de pedidos e é um item numerado, trate como "item_pedido"
    if em_pedidos and re.match(r'^\s*\d+[\.\)]\s+', texto_limpo):
        return 'item_pedido', False, 'justify'  # Justificar itens de pedidos
   

    # Itens Doc. - detecção robusta
    if re.match(r'^\s*Doc\.\s*\d+', texto_limpo):
        return 'item_doc', False, 'left'

    # Cabeçalho judicial (EXMO ou EXCELENTÍSSIMO(A))
    if re.match(r'^(EXMO|EXCELENTÍSSIM[OA])\b', texto_limpo, re.IGNORECASE):
        return 'cabecalho', True, 'center'

    # ETAPA 2: VERIFICAÇÃO DE CITAÇÕES APRIMORADA
    # Lista completa de aspas (incluindo Unicode e casos especiais)
    aspas = [
        # Aspas retas
        '"', "'",
        # Aspas curvas
        '“', '”', '‘', '’', 
        # Aspas angulares
        '«', '»', '‹', '›',
        # Aspas duplas baixas
        '„', '‟',
        # Aspas orientais
        '「', '」', '『', '』'
    ]

    # Verificar qualquer tipo de aspas no texto
    if any(aspas in texto_limpo for aspas in aspas):
        return 'citacao', False, 'justify'
    
    # ETAPA 3: Artigos de lei e referências jurídicas específicas
    # Estas são citações mesmo sem aspas, por serem referências técnicas
    if (re.match(r'Art\.\s*\d+', texto_limpo) or
        re.match(r'§\s*\d+', texto_limpo) or
        re.search(r'inciso\s+[IVX]+', texto_limpo) or
        re.search(r'alínea\s+[a-z]', texto_limpo)):
        return 'citacao', False, 'justify'
    
    # ETAPA 4: Verificações de marcadores e listas
    
    # Marcadores de lista com espaços ou não
    if re.match(r'^\s*[•▪■□◊○●◉◎◌◦⦿⦾]+\s+', texto_limpo):
        return 'subsecao', True, 'left'
    
    # Listas numeradas (números seguidos de ponto ou parêntese)
    if re.match(r'^\s*\d+[\.\)]\s+', texto_limpo):
        return 'lista', False, 'left'
    
    # Listas com letras (a., b., etc.)
    if re.match(r'^\s*[a-z][\.\)]\s+', texto_limpo):
        return 'lista', False, 'left'
    
    # ETAPA 5: Verificações de conteúdo específico
    
    # Seções principais - padrão romano e estrutura específica
    if re.match(r'^[IVX]+[\s]*[.–—\-]+[\s]*(DOS?|DAS?)[\s]+[A-ZÀÁÂÃÉÊÍÓÔÕÚÇ\s]+$', texto_limpo):
        return 'secao_principal', True, 'left'
    
    # ETAPA 6: Verificações de conteúdo baseadas em palavras-chave
    
    # Título da ação - critérios mais específicos
    # Verificar se contém "AÇÃO DE" e está em maiúsculas sem ser parte de outro item
    if ('AÇÃO DE' in texto_limpo.upper() and 
        len(texto_limpo) < 150 and 
        texto_limpo.upper().count(' ') >= 2 and  # Ter pelo menos 2 espaços (3 palavras)
        not any(marcador in texto_limpo for marcador in ['•', '▪', '-', '*']) and  # Não conter marcadores
        not re.match(r'^\s*\d+[\.\)]', texto_limpo)):  # Não ser item numerado
        return 'titulo_acao', True, 'center'
    
    # ETAPA 7: Verificações de outras características de formatação
    
    # Outros tipos de listas com marcadores diversos
    if re.match(r'^\s*[\-–—*+]\s+', texto_limpo):
        return 'lista', False, 'left'
    
    # ETAPA 8: Verificações baseadas em análise contextual
    
    # Textos que parecem títulos são formatados como normal para compatibilidade
    words = texto_limpo.split()
    if (2 <= len(words) <= 7 and 
        all(w[0].isupper() for w in words if len(w) > 3) and 
        not texto_limpo.endswith('.') and
        len(texto_limpo) < 50):
        return 'normal', True, 'left'  # Mantém como normal mas com negrito
    
    # ETAPA 9: Classificação padrão
    
    # Parágrafo normal
    return 'normal', False, 'justify'


def formatar_documento(doc_entrada, doc_saida_path, logo_path=None, debug_mode=False):
    # Lista para armazenar informações de depuração
    debug_info = []
    
    # Criar novo documento
    doc_novo = Document()

    # Configurar margens para o corpo do documento
    sections = doc_novo.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Adicionar cabeçalho com logo
    criar_cabecalho(doc_novo, logo_path)
    
    em_pedidos = False  # Desativa a seção de pedidos ao encontrar parágrafo vazio
    
    # Processar cada parágrafo do documento original
    for i, para in enumerate(doc_entrada.paragraphs):
        texto = para.text.strip()
         # Adicione esta verificação ANTES de pular parágrafos vazios
        if texto == '' and em_pedidos:
            em_pedidos = False  # Desativa a seção de pedidos ao encontrar parágrafo vazio

        if not texto:  # Pular parágrafos vazios mas adicionar espaço
            doc_novo.add_paragraph()
            continue
    

        # Detectar tipo de parágrafo
        tipo, negrito, alinhamento = detectar_tipo_paragrafo(texto, em_pedidos)
        
        # Armazenar informações para depuração
        if debug_mode:
            debug_info.append({
                "index": i,
                "texto": texto[:50] + "..." if len(texto) > 50 else texto,
                "tipo_detectado": tipo,
                "negrito": negrito,
                "alinhamento": alinhamento,
                "em_pedidos": em_pedidos,  # Novo campo
                "alinhamento_aplicado": 'justify' if em_pedidos else alinhamento  # Novo campo
            })

        # Criar novo parágrafo
        p = doc_novo.add_paragraph()
        run = p.add_run(texto)

        # Ativar modo Pedidos quando detectado
        if tipo == 'secao_pedidos':
            em_pedidos = True
        # Aplicar formatação baseada no tipo
        if tipo == 'cabecalho':
           aplicar_formatacao_paragrafo(p, alinhamento='center', negrito=True,
                               tamanho_fonte=12, espacamento_antes=0,
                               espacamento_depois=40, recuo_primeira_linha=False)
        
        elif tipo == 'titulo_acao':
             aplicar_formatacao_paragrafo(p, alinhamento='center', negrito=True,
                               tamanho_fonte=12, espacamento_antes=30,
                               espacamento_depois=24,
                               cor_texto=FORMATO_CONFIG['cor_titulo'],
                               recuo_primeira_linha=False)
        
        elif tipo == 'secao_principal':
             aplicar_formatacao_paragrafo(p, alinhamento='left', negrito=True,
                               tamanho_fonte=12, espacamento_antes=12,
                               espacamento_depois=6,
                               cor_texto=FORMATO_CONFIG['cor_secao'],
                               recuo_primeira_linha=False)
            # Adicionar linha horizontal cinza
             adicionar_linha_horizontal(p, FORMATO_CONFIG['cor_linha'])
        
        # Para itens de documentos específicos (Doc. X)
        elif tipo == 'item_doc':
            aplicar_formatacao_paragrafo(p, alinhamento='left', negrito=False,
                              tamanho_fonte=12, espacamento_antes=6,
                              espacamento_depois=6, recuo_lista=True)
        
        elif tipo == 'subsecao':
            aplicar_formatacao_paragrafo(p, alinhamento='left', negrito=True,
                              tamanho_fonte=12, espacamento_antes=6,
                              espacamento_depois=6, recuo_lista=True)
        
        elif tipo == 'citacao':
            aplicar_formatacao_paragrafo(p, alinhamento='justify', negrito=False,
                              italico=True, tamanho_fonte=11,  # Fonte menor e itálico
                              espacamento_antes=6, espacamento_depois=6,
                              recuo_lista=True)  # Com recuo para destacar
        
        elif tipo == 'lista':
            aplicar_formatacao_paragrafo(p, alinhamento='left', negrito=False,
                              tamanho_fonte=12, espacamento_antes=3,
                              espacamento_depois=3, recuo_lista=True)
                # Substitua o bloco atual por:
        elif tipo == 'secao_pedidos':
            aplicar_formatacao_paragrafo(p, 
                alinhamento='center',
                negrito=True,
                tamanho_fonte=12,
                espacamento_antes=24,
                espacamento_depois=12,
                recuo_primeira_linha=False
            )
            adicionar_linha_horizontal(p, FORMATO_CONFIG['cor_linha'])
            em_pedidos = True  # Ativa o modo pedidos
        
        # E modifique a formatação do conteúdo dos pedidos:
        elif tipo == 'item_pedido':
            aplicar_formatacao_paragrafo(p, 
                alinhamento='justify',  # Justificado
                negrito=False,
                tamanho_fonte=12,
                espacamento_antes=6,
                espacamento_depois=6,
                recuo_lista=True  # Manter o recuo de lista
            )
        elif em_pedidos:
            aplicar_formatacao_paragrafo(p, 
                alinhamento='justify',  # Alterado para justificado
                negrito=False,
                tamanho_fonte=12,
                espacamento_antes=6,
                espacamento_depois=6,
                recuo_primeira_linha=True
            )
      
        
        else:  # normal
            atual_alinhamento = 'justify' 
            aplicar_formatacao_paragrafo(p, alinhamento=atual_alinhamento, negrito=False,
                               tamanho_fonte=12, espacamento_antes=6,
                               espacamento_depois=6)
        

    # Processar tabelas do documento original
    for table in doc_entrada.tables:
        # Criar nova tabela com mesma estrutura
        nova_tabela = doc_novo.add_table(rows=len(table.rows), cols=len(table.columns))
        nova_tabela.style = 'Light Grid Accent 1'

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                nova_tabela.rows[i].cells[j].text = cell.text
                # Formatar primeira linha como cabeçalho
                if i == 0:
                    for paragraph in nova_tabela.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)

        # Adicionar espaço após tabela
        doc_novo.add_paragraph()

    # Configuração especial para o rodapé de página inteira
    # Obter a última seção do documento (onde o rodapé será aplicado)
    last_section = doc_novo.sections[-1]

    # Fazer uma cópia das margens originais do documento
    original_left = last_section.left_margin
    original_right = last_section.right_margin

    # Adicionar rodapé com as margens padrão
    criar_rodape(doc_novo, RODAPE_CONFIG)

    # Modificar as propriedades do rodapé para ocupar a largura total
    footer = last_section.footer

    # Aplicar estilo especial ao parágrafo do rodapé
    for p in footer.paragraphs:
        if p.text.strip():  # Se não estiver vazio
            # Estender o parágrafo além das margens
            p_format = p.paragraph_format

            # Usar valores XML diretos para estender além das margens
            p_element = p._element.get_or_add_pPr()

            # Adicionar configuração de moldura para estender além das margens com altura fixa
            from docx.oxml import parse_xml
            frame_xml = parse_xml(
            '<w:framePr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'w:w="13000" w:h="2500" w:wrap="around" w:vAnchor="page" w:hAnchor="page" w:xAlign="center" />'
             )
            p_element.append(frame_xml)

            # Definir espaçamento interno para o rodapé ter altura de 2.4cm
            p.paragraph_format.space_before = Pt(50)  # Aproximadamente 1.2 cm
            p.paragraph_format.space_after = Pt(50)   # Aproximadamente 1.2 cm

    # Salvar documento
    doc_novo.save(doc_saida_path)
    
    if debug_mode:
        return doc_saida_path, debug_info
    else:
        return doc_saida_path


def criar_arquivo_zip(arquivos):
    """Cria um arquivo ZIP contendo todos os arquivos processados"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for arquivo in arquivos:
            zip_file.write(arquivo, os.path.basename(arquivo))
    
    zip_buffer.seek(0)
    return zip_buffer


def main():
    st.set_page_config(
        page_title="Formatador Jurídico ICA Advocacia", 
        page_icon="📄",
        layout="wide"
    )

    # Inicializar variável de estado para modo de depuração
    if 'debug_mode' not in st.session_state:
        st.session_state.debug_mode = False

    st.title("📄 Formatador Jurídico ICA Advocacia")
    st.write("Ferramenta para formatação automática de documentos jurídicos.")
    
    # Adicionar espaço e linha separadora para melhor organização visual
    st.markdown("---")

    # Sidebar para configurações
    with st.sidebar:
        st.header("⚙️ Configurações")
        st.write("**Opções de formatação:**")
        
        # Guardar o logo padrão
        save_logo = st.checkbox("Salvar logo para uso futuro", value=True)
        
        # Opção de depuração
        st.session_state.debug_mode = st.checkbox("Modo de depuração", value=False)
        if st.session_state.debug_mode:
            st.info("O modo de depuração mostrará informações detalhadas sobre a formatação.")
    
    # ETAPA 1: Upload do Logo - Em seção separada e bem visível
    st.header("1️⃣ Upload do Logo ICA")
    
    # Criar caixa para o upload do logo
    logo_container = st.container()
    with logo_container:
        logo_col1, logo_col2 = st.columns([2, 1])
        
        with logo_col1:
            logo_file = st.file_uploader("Faça upload do logo ICA (arquivo PNG, JPG)", 
                                        type=["png", "jpg", "jpeg"], 
                                        key="logo")

        with logo_col2:
            # Verificar se existe logo salvo em cache
            if 'logo_cache' in st.session_state and logo_file is None:
                logo_path = st.session_state.logo_cache
                st.success("✅ Logo salvo em uso")
                st.image(logo_path, width=150)
            elif logo_file is not None:
                # Salvar o logo em um arquivo temporário
                temp_logo = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                temp_logo.write(logo_file.getvalue())
                logo_path = temp_logo.name

                # Salvar em cache se a opção estiver marcada
                if save_logo:
                    st.session_state.logo_cache = logo_path
                
                # Mostrar confirmação e preview
                st.success("✅ Logo carregado!")
                st.image(logo_file, width=150)
            else:
                logo_path = None
                st.info("Logo não carregado. O cabeçalho será criado sem logo.")
    
    # Adicionar espaço e linha separadora para melhor organização visual
    st.markdown("---")
    
    # ETAPA 2: Upload dos Documentos Word - Em seção separada
    st.header("2️⃣ Upload dos Documentos Word")
    
    doc_container = st.container()
    with doc_container:
        uploaded_files = st.file_uploader(
            "Selecione os documentos Word (.docx) para formatação",
            type=["docx"], 
            accept_multiple_files=True,
            key="docs"
        )

        if not uploaded_files:
            st.warning("⚠️ Nenhum documento carregado ainda.")
        else:
            st.success(f"✅ {len(uploaded_files)} documento(s) carregado(s)")
            
            # Mostrar lista de arquivos carregados
            with st.expander(f"📋 Ver arquivos carregados ({len(uploaded_files)})"):
                for i, doc_file in enumerate(uploaded_files):
                    st.write(f"{i+1}. {doc_file.name}")
    
    # Adicionar espaço e linha separadora para melhor organização visual
    st.markdown("---")

    # ETAPA 3: Formatação - Em seção separada
    st.header("3️⃣ Formatação")
    
    # Botão para formatar documentos
    format_button = st.button(
        "Formatar Documentos", 
        disabled=(len(uploaded_files) == 0 or (logo_path is None and 'logo_cache' not in st.session_state)),
        type="primary",
        use_container_width=True
    )

    # Processamento dos documentos quando o botão é pressionado
    if format_button and len(uploaded_files) > 0:
        # Usar o logo em cache se disponível e nenhum foi carregado
        if logo_path is None and 'logo_cache' in st.session_state:
            logo_path = st.session_state.logo_cache
            
        # Criar área de status e progresso
        progress_container = st.container()
        
        with progress_container:
            status_text = st.empty()
            progress_bar = st.progress(0)
            
            # Criar pasta de saída temporária
            temp_dir = tempfile.mkdtemp()
            
            # Indicador de arquivos processados
            processed_count = st.empty()
            processed_count.info("Preparando processamento...")
            
            # Processar cada arquivo
            arquivos_processados = []
            errors = []
            
            for i, doc_file in enumerate(uploaded_files):
                try:
                    # Atualizar status
                    status_text.info(f"⏳ Processando: {doc_file.name} ({i+1}/{len(uploaded_files)})")
                    
                    # Salvar arquivo temporariamente
                    input_path = os.path.join(temp_dir, doc_file.name)
                    with open(input_path, "wb") as f:
                        f.write(doc_file.getvalue())
                    
                    # Gerar nome de saída
                    nome_base = os.path.splitext(doc_file.name)[0]
                    output_path = os.path.join(temp_dir, f"{nome_base}_FORMATADO.docx")
                    
                    # Formatar documento (agora com opção de debug)
                    doc = Document(input_path)
                    
                    # Chama a função com o modo de depuração
                    resultado = formatar_documento(doc, output_path, logo_path, st.session_state.debug_mode)
                    
                    if st.session_state.debug_mode:
                        output_path, debug_info = resultado
                        st.session_state[f'debug_info_{i}'] = debug_info
                    else:
                        output_path = resultado
                        
                    arquivos_processados.append(output_path)
                    
                    # Atualizar progresso
                    progress = int(((i + 1) / len(uploaded_files)) * 100)
                    progress_bar.progress(progress)
                    processed_count.info(f"✅ Processados: {i+1}/{len(uploaded_files)} documentos")
                    
                except Exception as e:
                    errors.append((doc_file.name, str(e)))
            
            # Finalizar processamento
            if arquivos_processados:
                status_text.success(f"✅ Processamento concluído! {len(arquivos_processados)} documento(s) formatado(s).")
                progress_bar.progress(100)
                
                # Mostrar erros, se houver
                if errors:
                    with st.expander(f"⚠️ Erros ({len(errors)})"):
                        for file_name, error_msg in errors:
                            st.error(f"Arquivo: {file_name} - Erro: {error_msg}")
                
                # Exibir informações de depuração se ativado
                if st.session_state.debug_mode:
                    st.markdown("---")
                    st.header("🔍 Informações de Depuração")
                    
                    for i, doc_file in enumerate(uploaded_files):
                        debug_key = f'debug_info_{i}'
                        if debug_key in st.session_state:
                            with st.expander(f"Debug: {doc_file.name}"):
                                # Criar uma tabela com as informações de depuração
                                st.write("### Análise de Parágrafos")
                                
                                # Tabela de depuração
                                debug_data = st.session_state[debug_key]
                                
                                # Criar tabela
                                st.table([{
                                    "#": item["index"],
                                    "Texto": item["texto"],
                                    "Tipo": item["tipo_detectado"],
                                    "Negrito": "Sim" if item["negrito"] else "Não",
                                    "Alinhamento": item["alinhamento"]
                                } for item in debug_data])
                                
                                # Destacar possíveis problemas
                                st.subheader("Possíveis problemas detectados")
                                
                                problemas = []
                                for item in debug_data:
                                    if "Doc." in item["texto"] and item["tipo_detectado"] != "item_doc":
                                        problemas.append({
                                            "Parágrafo": item["index"],
                                            "Texto": item["texto"],
                                            "Problema": f"Item Doc. detectado como '{item['tipo_detectado']}'"
                                        })
                                    if "•" in item["texto"] and item["tipo_detectado"] != "subsecao" and item["tipo_detectado"] != "lista":
                                        problemas.append({
                                            "Parágrafo": item["index"],
                                            "Texto": item["texto"],
                                            "Problema": f"Marcador • detectado como '{item['tipo_detectado']}'"
                                        })
                                    if 'PEDIDOS' in item["texto"].upper() or 'POR TUDO ISSO' in item["texto"].upper():
                                        # Verificar se detectou corretamente como seção de pedidos
                                        if item["tipo_detectado"] != "secao_pedidos":
                                            problemas.append({
                                                "Parágrafo": item["index"],
                                                "Texto": item["texto"],
                                                "Problema": f"Seção de Pedidos detectada como '{item['tipo_detectado']}' (deveria ser 'secao_pedidos')"
                                            })
                                        # Verificar alinhamento
                                        elif item["alinhamento"] != "justify":
                                            problemas.append({
                                                "Parágrafo": item["index"],
                                                "Texto": item["texto"],
                                                "Problema": f"Alinhamento incorreto na seção de Pedidos: '{item['alinhamento']}' (deveria ser 'justify')"
                                            })
                                        # Verificar parágrafos subsequentes aos Pedidos
                                        if item["tipo_detectado"] == "secao_pedidos":
                                            # Encontrar parágrafos seguintes até próxima seção
                                            proximos_paragrafos = [p for p in debug_data if p["index"] > item["index"] and p["texto"].strip() != ""]
                                            for p in proximos_paragrafos[:5]:  # Verificar os 5 próximos parágrafos
                                                if p["alinhamento"] != "justify":
                                                    problemas.append({
                                                        "Parágrafo": p["index"],
                                                        "Texto": p["texto"],
                                                        "Problema": f"Parágrafo após Pedidos com alinhamento '{p['alinhamento']}' (deveria ser 'justify')"
                                                    })
                                if problemas:
                                    st.table(problemas)
                                else:
                                    st.success("Nenhum problema evidente detectado.")
                    
                # Adicionar espaço e linha separadora
                st.markdown("---")
                
                # ETAPA 4: Download - Em seção separada
                st.header("4️⃣ Download dos Documentos Formatados")
                
                # Criar nome para o arquivo ZIP
                data_atual = datetime.now().strftime("%Y%m%d")
                zip_filename = f"Documentos_Formatados_ICA_{data_atual}.zip"
                
                # Botão grande e destacado para download em lote
                if len(arquivos_processados) > 1:
                    st.subheader("📦 Download de todos os arquivos")
                    zip_data = criar_arquivo_zip(arquivos_processados)
                    
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.download_button(
                            label="⬇️ BAIXAR TODOS OS DOCUMENTOS DE UMA VEZ",
                            data=zip_data,
                            file_name=zip_filename,
                            mime="application/zip",
                            use_container_width=True,
                            type="primary",
                        )
                    
                    with col2:
                        st.info(f"{len(arquivos_processados)} arquivos no ZIP")
                    
                    # Linha separadora
                    st.markdown("---")
                
                # Downloads individuais em seção separada
                st.subheader("📄 Downloads individuais")
                
                # Criar grid mais organizado para os arquivos individuais
                num_cols = 2  # Reduzido para 2 colunas para melhor espaçamento
                file_cols = st.columns(num_cols)
                
                for i, file_path in enumerate(arquivos_processados):
                    col_idx = i % num_cols
                    with file_cols[col_idx]:
                        file_name = os.path.basename(file_path)
                        with open(file_path, "rb") as file:
                            st.download_button(
                                label=f"⬇️ {file_name}",
                                data=file,
                                file_name=file_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_{i}",
                                use_container_width=True
                            )
                        # Adicionar espaço entre botões
                        st.write("")
            else:
                st.error("❌ Nenhum documento foi processado com sucesso.")

    # Informações adicionais em rodapé
    st.markdown("---")
    with st.expander("ℹ️ Informações sobre o Formatador"):
        st.markdown("""
        ### Sobre a Formatação
        
        Este aplicativo formata documentos jurídicos seguindo o padrão visual do escritório ICA Advocacia:
        
        - **Cabeçalho**: Logo centralizado
        - **Títulos**: Em azul com formatação adequada
        - **Seções principais**: Com linha horizontal
        - **Rodapé**: Em azul escuro com informações do escritório
        
        ### Dicas de Uso
        
        - Você pode processar vários documentos de uma vez
        - Para melhor organização, use nomes descritivos para seus arquivos
        - O logo será mantido para uso futuro se você marcar a opção na barra lateral
        - Ative o modo de depuração na barra lateral para ajudar a identificar problemas
        """)
    
    # Adicionar rodapé discreto da aplicação
    st.markdown("---")
    st.markdown("<div style='text-align: center; color: gray; font-size: 0.8em;'>Formatador Jurídico ICA Advocacia - Versão 1.0</div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
